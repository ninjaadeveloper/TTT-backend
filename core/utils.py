import os
import re
import requests
from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT


# --- Font setup for PDF (Unicode safe) ---
FONT_DIR = "storage/fonts"
FONT_PATH = os.path.join(FONT_DIR, "NotoSans.ttf")


def ensure_font():
    """Download NotoSans variable font from Google Fonts repo if not available."""
    os.makedirs(FONT_DIR, exist_ok=True)
    if not os.path.exists(FONT_PATH):
        print("🔽 Downloading NotoSans.ttf ...")
        url = ("https://github.com/google/fonts/raw/main/ofl/notosans/"
               "NotoSans%5Bwdth,wght%5D.ttf")
        r = requests.get(url, timeout=30)
        r.raise_for_status()
        with open(FONT_PATH, "wb") as f:
            f.write(r.content)
        print("✅ NotoSans font downloaded!")


# --- Export Notes to PDF (Enhanced) ---
def export_to_pdf(notes_text, output_path):
    """
    Enhanced PDF export – supports ## headings, - bullets,
    numbered lists, and normal text.
    """
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    ensure_font()  # Ensure NotoSans font is available
    pdfmetrics.registerFont(TTFont("NotoSans", FONT_PATH))

    doc = SimpleDocTemplate(
        output_path,
        pagesize=A4,
        rightMargin=50,
        leftMargin=50,
        topMargin=50,
        bottomMargin=50,
    )

    styles = getSampleStyleSheet()
    story = []

    heading_style = ParagraphStyle(
        "Heading",
        parent=styles["Heading2"],
        fontName="NotoSans",
        fontSize=14,
        leading=18,
        spaceAfter=10,
        textColor=colors.HexColor("#222222"),
    )

    body_style = ParagraphStyle(
        "Body",
        parent=styles["Normal"],
        fontName="NotoSans",
        fontSize=12,
        leading=16,
        spaceAfter=8,
    )

    bullet_style = ParagraphStyle(
        "Bullet",
        parent=styles["Normal"],
        fontName="NotoSans",
        fontSize=12,
        leading=16,
        leftIndent=20,
        bulletIndent=10,
        spaceAfter=6,
    )

    lines = notes_text.splitlines()

    for line in lines:
        line = line.strip()
        if not line:
            story.append(Spacer(1, 8))
            continue

        if line.startswith("##"):
            story.append(Paragraph(line.replace("##", "").strip(),
                                   heading_style))
        elif re.match(r"^[-*]\s+", line):
            story.append(Paragraph("• " + line[2:].strip(), bullet_style))
        elif re.match(r"^\d+\.\s+", line):
            story.append(Paragraph(line.strip(), bullet_style))
        else:
            story.append(Paragraph(line, body_style))

    doc.build(story)
    return output_path


# --- Export Notes to DOCX (Enhanced) ---
def export_to_docx(notes_text, output_path):
    """
    Enhanced DOCX export – supports markdown-like structure (##, -, 1.)
    """
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc = Document()

    if isinstance(notes_text, str):
        lines = notes_text.split("\n")
    else:
        lines = list(notes_text)

    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph("")  # blank line for spacing
            continue

        if line.startswith("##"):
            doc.add_heading(line.replace("##", "").strip(), level=2)
        elif re.match(r"^[-*]\s+", line):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif re.match(r"^\d+\.\s+", line):
            doc.add_paragraph(line.strip(), style="List Number")
        else:
            doc.add_paragraph(line)

    doc.save(output_path)
    return output_path


# --- Extract Audio from Video ---
def extract_audio_from_video(video_path, output_path, duration=120):
    """
    Extract first `duration` seconds of audio from a video (e.g. MP4).
    Save as .mp3 file.
    """
    from moviepy.editor import VideoFileClip

    if not os.path.exists(video_path):
        raise FileNotFoundError(f"Video not found: {video_path}")

    clip = VideoFileClip(video_path)
    audio_clip = None
    try:
        if clip.audio is None:
            raise RuntimeError("No audio track found in video!")

        audio_clip = clip.audio.subclip(0, duration)
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        audio_clip.write_audiofile(output_path, codec="mp3")
    finally:
        if audio_clip:
            try:
                audio_clip.close()
            except Exception:
                pass
        if clip:
            try:
                clip.close()
            except Exception:
                pass

    return output_path


# --- Translation helper (uses googletrans, fallback to identity) ---
def translate_text(text, src="auto", target="en"):
    """
    Translate `text` to target language. Uses googletrans if installed.
    Returns translated text or original text on failure.
    """
    try:
        from googletrans import Translator

        translator = Translator()
        if len(text) <= 4500:
            return translator.translate(text, src=src, dest=target).text
        out = []
        start = 0
        while start < len(text):
            chunk = text[start:start + 4500]
            out.append(translator.translate(chunk, src=src, dest=target).text)
            start += 4500
        return "\n".join(out)
    except Exception:
        return text


# --- Token/text optimizer (heuristic) ---
def optimize_for_tokens(text, max_tokens=3000):
    """
    Heuristic to reduce text length to approximately max_tokens.
    Approx tokens = chars/4 (simple estimate).
    Tries to cut at sentence boundary.
    """
    if not text:
        return text
    approx_tokens = len(text) / 4.0
    if approx_tokens <= max_tokens:
        return text

    ratio = max_tokens / approx_tokens
    max_chars = max(200, int(len(text) * ratio))
    cut = text[:max_chars]
    last_dot = max(cut.rfind("."), cut.rfind("!\n"), cut.rfind("?\n"))
    if last_dot and last_dot > int(0.5 * max_chars):
        return cut[: last_dot + 1]
    last_nl = cut.rfind("\n")
    if last_nl and last_nl > 0:
        return cut[:last_nl]
    return cut
